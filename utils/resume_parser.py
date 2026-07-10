import fitz  # PyMuPDF
import docx


def extract_text_from_pdf(file_path):
    text_parts = []
    with fitz.open(file_path) as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)


def extract_text_from_docx(file_path):
    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs]

    # tables sometimes hold skills/education in resume templates — grab those too
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def extract_text(file_path, extension):
    extension = extension.lower()
    if extension == "pdf":
        return extract_text_from_pdf(file_path)
    if extension == "docx":
        return extract_text_from_docx(file_path)
    raise ValueError(f"Unsupported file type: {extension}")
